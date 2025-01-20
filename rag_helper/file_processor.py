import os
import json
import csv
import fitz  # PyMuPDF
import docx
import pickle
from bs4 import BeautifulSoup
from pathlib import Path
from typing import List, Optional, Dict
from tqdm import tqdm
import logging

class FileChunker:
    """Utility to split various document types into smaller text chunks."""
    
    def __init__(
        self, 
        input_file: str, 
        chunk_size: int = 500, 
        method: str = "sentence",
        output_dir: Optional[str] = None
    ):
        """
        Initialize the file chunker.
        
        Args:
            input_file: Path to the file to chunk
            chunk_size: Approximate number of words per chunk
            method: Chunking method - "sentence" or "token"
            output_dir: Custom output directory (optional)
        """
        self.input_file = Path(input_file)
        self.chunk_size = chunk_size
        self.method = method
        
        # Set up logging
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        
        # Validate input file
        if not self.input_file.exists():
            raise FileNotFoundError(f"File not found: {input_file}")
            
        if method not in ["sentence", "token"]:
            raise ValueError("Method must be either 'sentence' or 'token'")
            
        # Set up output directory
        if output_dir:
            self.output_dir = Path(output_dir)
        else:
            # Create default output directory next to the input file
            default_dir = self.input_file.parent / "chunks" / self.input_file.stem
            self.output_dir = default_dir
        
        # Create output directory if it doesn't exist
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Save metadata
        self.metadata = {
            "original_file": str(self.input_file),
            "chunk_size": chunk_size,
            "method": method,
            "output_dir": str(self.output_dir)
        }
    
    def read_file(self) -> str:
        """Read content from various file formats."""
        suffix = self.input_file.suffix.lower()
        
        try:
            if suffix == '.txt':
                return self._read_text_file()
            elif suffix == '.pdf':
                return self._read_pdf_file()
            elif suffix == '.docx':
                return self._read_docx_file()
            elif suffix == '.json':
                return self._read_json_file()
            elif suffix == '.csv':
                return self._read_csv_file()
            elif suffix in ['.md', '.markdown']:
                return self._read_text_file()
            elif suffix == '.html':
                return self._read_html_file()
            elif suffix == '.xml':
                return self._read_xml_file()
            elif suffix == '.rtf':
                return self._read_rtf_file()
            elif suffix == '.pkl':
                return self._read_pickle_file()
            else:
                raise ValueError(f"Unsupported file format: {suffix}")
        except Exception as e:
            logging.error(f"Error reading {self.input_file}: {str(e)}")
            raise
    
    def _read_text_file(self) -> str:
        """Read content from text files (txt, md)."""
        with open(self.input_file, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _read_pdf_file(self) -> str:
        """Read content from PDF files using PyMuPDF."""
        text = []
        doc = fitz.open(self.input_file)
        for page_num in range(doc.page_count):
            text.append(doc[page_num].get_text())
        doc.close()
        return '\n'.join(text)
    
    def _read_docx_file(self) -> str:
        """Read content from Word documents."""
        doc = docx.Document(self.input_file)
        text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text.append(para.text)
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
                    
        return '\n'.join(text)
    
    def _read_json_file(self) -> str:
        """Read content from JSON files with nested structure support."""
        def extract_text(obj) -> List[str]:
            if isinstance(obj, str):
                return [obj]
            elif isinstance(obj, (int, float, bool)):
                return [str(obj)]
            elif isinstance(obj, dict):
                return [item for v in obj.values() for item in extract_text(v)]
            elif isinstance(obj, list):
                return [item for v in obj for item in extract_text(v)]
            return []

        with open(self.input_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return ' '.join(extract_text(data))
    
    def _read_csv_file(self) -> str:
        """Read content from CSV files with dialect detection."""
        text = []
        with open(self.input_file, 'r', encoding='utf-8', newline='') as f:
            sample = f.read(1024)
            f.seek(0)
            dialect = csv.Sniffer().sniff(sample)
            reader = csv.reader(f, dialect)
            for row in reader:
                text.append(' '.join(str(cell) for cell in row))
        return '\n'.join(text)
    
    def _read_html_file(self) -> str:
        """Read content from HTML files."""
        with open(self.input_file, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f, 'html.parser')
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            return soup.get_text(separator='\n')
    
    def _read_xml_file(self) -> str:
        """Read content from XML files."""
        with open(self.input_file, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f, 'xml')
            return soup.get_text(separator='\n')
    
    def _read_rtf_file(self) -> str:
        """Read content from RTF files."""
        # Simple RTF to text conversion
        with open(self.input_file, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
            # Remove RTF formatting
            text = []
            in_control = False
            for char in content:
                if char == '{' or char == '}' or char == '\\':
                    in_control = True
                    continue
                if char.isspace():
                    in_control = False
                if not in_control and char.isprintable():
                    text.append(char)
            return ''.join(text)
    
    def _read_pickle_file(self) -> str:
        """Read content from pickle files."""
        with open(self.input_file, 'rb') as f:
            data = pickle.load(f)
            if isinstance(data, str):
                return data
            return str(data)
    
    def chunk_by_sentences(self, text: str) -> List[str]:
        """Split text into chunks at sentence boundaries."""
        sentences = text.replace('\n', ' ').split('. ')
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            sentence_length = len(sentence.split())
            if current_length + sentence_length <= self.chunk_size:
                current_chunk.append(sentence)
                current_length += sentence_length
            else:
                if current_chunk:
                    chunks.append('. '.join(current_chunk) + '.')
                current_chunk = [sentence]
                current_length = sentence_length
                
        if current_chunk:
            chunks.append('. '.join(current_chunk) + '.')
        return chunks
    
    def chunk_by_tokens(self, text: str) -> List[str]:
        """Split text into chunks of fixed word count."""
        words = text.replace('\n', ' ').split()
        chunks = []
        
        for i in range(0, len(words), self.chunk_size):
            chunk = ' '.join(words[i:i + self.chunk_size])
            if chunk:
                chunks.append(chunk)
        return chunks
    
    def save_chunks(self) -> Dict[str, str]:
        """
        Process the input file and save chunks.
        
        Returns:
            Dict[str, str]: Mapping of chunk IDs to their file paths
        """
        # Read input file
        logging.info(f"Reading {self.input_file}...")
        text = self.read_file()
        
        # Generate chunks
        logging.info("Generating chunks...")
        if self.method == "sentence":
            chunks = self.chunk_by_sentences(text)
        else:
            chunks = self.chunk_by_tokens(text)
        
        # Save chunks with progress bar
        chunk_mapping = {}
        for i, chunk in enumerate(tqdm(chunks, desc="Saving chunks")):
            chunk_id = f"chunk_{i+1}"
            chunk_path = self.output_dir / f"{chunk_id}.txt"
            chunk_path.write_text(chunk, encoding='utf-8')
            chunk_mapping[chunk_id] = str(chunk_path)
        
        # Save metadata
        self.metadata["total_chunks"] = len(chunks)
        metadata_path = self.output_dir / "metadata.json"
        with open(metadata_path, 'w', encoding='utf-8') as f:
            json.dump(self.metadata, f, indent=2)
        
        logging.info(f"Created {len(chunks)} chunks in {self.output_dir}")
        return chunk_mapping

